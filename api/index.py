from flask import Flask, request, send_file, jsonify
from pypdf import PdfReader
from docx import Document
import os

app = Flask(__name__)

# Vercel hanya mengizinkan tulis file di folder /tmp
UPLOAD_FOLDER = '/tmp'

@app.route('/api/convert', methods=['POST'])
def convert_file():
    # 1. Validasi File
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    # 2. Simpan file PDF sementara
    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    # Ganti ekstensi jadi .docx
    output_filename = os.path.splitext(file.filename)[0] + ".docx"
    output_path = os.path.join(UPLOAD_FOLDER, output_filename)
    
    file.save(input_path)

    try:
        # --- PROSES KONVERSI (Versi Ringan) ---
        # Baca PDF
        reader = PdfReader(input_path)
        
        # Buat File Word Baru
        doc = Document()
        doc.add_heading(f'Hasil Konversi: {file.filename}', 0)
        
        # Ekstrak teks per halaman
        text_found = False
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_found = True
                doc.add_paragraph(text)
                doc.add_page_break()
        
        if not text_found:
             doc.add_paragraph("[Sistem mendeteksi PDF ini berisi Gambar/Scan. Teks tidak dapat diekstrak dengan metode ringan.]")

        # Simpan Word
        doc.save(output_path)
        
        # 3. Kirim File ke User
        return send_file(output_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
    finally:
        # Bersihkan file sampah di /tmp
        if os.path.exists(input_path):
            os.remove(input_path)
        # Output path akan dihapus sistem nanti, atau bisa dihapus manual setelah return (membutuhkan fungsi 'after_request')
